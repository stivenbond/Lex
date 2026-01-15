from fpdf import FPDF
from docx import Document
import sqlite3

class ExportService:
    @staticmethod
    def get_entry_data(entry_id):
        conn = sqlite3.connect("lex_database.db")
        cursor = conn.cursor()
        
        # Get entry basic info
        cursor.execute("""
            SELECT e.id, e.teaching_date, e.is_reference, e.reference_entry_id, 
                   t.name as topic_name, c.name as cohort_name, tr.name as teacher_name
            FROM diary_entries e
            JOIN topics t ON e.topic_id = t.id
            JOIN cohorts c ON e.cohort_id = c.id
            JOIN teachers tr ON e.teacher_id = tr.id
            WHERE e.id = ?
        """, (entry_id,))
        entry = cursor.fetchone()
        
        if not entry:
            conn.close()
            return None

        entry_dict = {
            "id": entry[0],
            "date": entry[1],
            "is_reference": entry[2],
            "ref_id": entry[3],
            "topic": entry[4],
            "cohort": entry[5],
            "teacher": entry[6],
            "values": []
        }

        if entry_dict["is_reference"]:
            # Get reference data
            cursor.execute("""
                SELECT e.teaching_date, c.name
                FROM diary_entries e
                JOIN cohorts c ON e.cohort_id = c.id
                WHERE e.id = ?
            """, (entry_dict["ref_id"],))
            ref_data = cursor.fetchone()
            entry_dict["reference_info"] = {
                "date": ref_data[0],
                "cohort": ref_data[1]
            }
        else:
            # Get field values
            cursor.execute("""
                SELECT f.label, v.value
                FROM diary_values v
                JOIN diary_fields f ON v.field_id = f.id
                WHERE v.entry_id = ?
            """, (entry_id,))
            entry_dict["values"] = cursor.fetchall()

        conn.close()
        return entry_dict

    @staticmethod
    def export_to_pdf(entry_id, output_path):
        data = ExportService.get_entry_data(entry_id)
        if not data: return

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Lesson Diary Entry", ln=True, align='C')
        
        pdf.set_font("Arial", size=12)
        pdf.ln(10)
        pdf.cell(200, 10, txt=f"Teacher: {data['teacher']}", ln=True)
        pdf.cell(200, 10, txt=f"Topic: {data['topic']}", ln=True)
        pdf.cell(200, 10, txt=f"Cohort: {data['cohort']}", ln=True)
        pdf.cell(200, 10, txt=f"Date: {data['date']}", ln=True)
        
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt="Content:", ln=True)
        pdf.set_font("Arial", size=11)
        
        if data["is_reference"]:
            ref = data["reference_info"]
            pdf.multi_cell(0, 10, txt=f"REFERENCED ENTRY: This lesson content is identical to the entry recorded on {ref['date']} for cohort {ref['cohort']}.")
        else:
            for label, value in data["values"]:
                pdf.set_font("Arial", 'B', 11)
                pdf.cell(0, 10, txt=f"{label}:", ln=True)
                pdf.set_font("Arial", size=11)
                pdf.multi_cell(0, 10, txt=value if value else "N/A")
                pdf.ln(2)

        pdf.output(output_path)

    @staticmethod
    def export_to_docx(entry_id, output_path):
        data = ExportService.get_entry_data(entry_id)
        if not data: return

        doc = Document()
        doc.add_heading('Lesson Diary Entry', 0)
        
        doc.add_paragraph(f"Teacher: {data['teacher']}")
        doc.add_paragraph(f"Topic: {data['topic']}")
        doc.add_paragraph(f"Cohort: {data['cohort']}")
        doc.add_paragraph(f"Date: {data['date']}")
        
        doc.add_heading('Content', level=1)
        
        if data["is_reference"]:
            ref = data["reference_info"]
            p = doc.add_paragraph()
            p.add_run("REFERENCED ENTRY: ").bold = True
            p.add_run(f"This lesson content is identical to the entry recorded on {ref['date']} for cohort {ref['cohort']}.")
        else:
            for label, value in data["values"]:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value if value else "N/A")

        doc.save(output_path)

    @staticmethod
    def export_whiteboard_to_pdf(topic_id, output_path):
        conn = sqlite3.connect("lex_database.db")
        cursor = conn.cursor()
        cursor.execute("SELECT content, x, y FROM text_blocks WHERE topic_id = ?", (topic_id,))
        blocks = cursor.fetchall()
        
        cursor.execute("SELECT name FROM topics WHERE id = ?", (topic_id,))
        topic_name = cursor.fetchone()[0]
        conn.close()

        pdf = FPDF(orientation='L', unit='pt', format='A4') # Landscape for whiteboard
        pdf.add_page()
        pdf.set_font("Arial", 'B', 20)
        pdf.text(40, 40, f"Whiteboard: {topic_name}")
        
        pdf.set_font("Arial", size=10)
        for content, x, y in blocks:
            # Scale positions if needed, for now assume they are in pt or similar
            # Ensure x, y are numbers
            px = x if x and x > 0 else 50
            py = y if y and y > 0 else 100
            
            # Draw a box
            pdf.set_fill_color(230, 230, 230)
            pdf.rect(px, py, 150, 40, 'F')
            pdf.set_xy(px + 5, py + 5)
            # Use multi_cell for wrapping text inside the box
            pdf.multi_cell(140, 10, txt=content[:100] + ("..." if len(content) > 100 else ""))

        pdf.output(output_path)
